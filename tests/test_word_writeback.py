"""Regression tests for character-level DOCX writeback and OOXML comments."""

import hashlib
import shutil
import subprocess
import tempfile
import unittest
import zipfile
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn as docx_qn
from docx.opc.constants import RELATIONSHIP_TYPE
from lxml import etree

from src.max_pipeline import _run_02_writeback
from src.writeback_engine import DocxAuditError, W14, W15, audit_docx, mcn, qn


def _add_hyperlink(paragraph, text, url):
    rel_id = paragraph.part.relate_to(
        url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(docx_qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _wrap_last_run_in_revision(paragraph, author="Human"):
    run = paragraph._p.findall(qn("r"))[-1]
    paragraph._p.remove(run)
    insertion = OxmlElement("w:ins")
    insertion.set(docx_qn("w:id"), "88")
    insertion.set(docx_qn("w:author"), author)
    insertion.set(docx_qn("w:date"), "2026-01-01T00:00:00Z")
    insertion.append(run)
    paragraph._p.append(insertion)


def _add_existing_comment_extension(path):
    """Add a pre-existing commentsExtended entry with state to preserve."""
    temp = path.with_suffix(".extended.docx")
    with zipfile.ZipFile(path, "r") as source:
        comments = etree.fromstring(source.read("word/comments.xml"))
        comment_para = next(comments.iter(qn("p")))
        comment_para.set(f"{{{W14}}}paraId", "ABCDEF01")
        comment_para.set(f"{{{W14}}}textId", "77777777")

        extended = etree.Element(f"{{{W15}}}commentsEx")
        old = etree.SubElement(extended, f"{{{W15}}}commentEx")
        old.set(f"{{{W15}}}paraId", "ABCDEF01")
        old.set(f"{{{W15}}}done", "1")

        content_types = etree.fromstring(source.read("[Content_Types].xml"))
        ct_ns = content_types.tag.split("}")[0].strip("{")
        override = etree.SubElement(content_types, f"{{{ct_ns}}}Override")
        override.set("PartName", "/word/commentsExtended.xml")
        override.set(
            "ContentType",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.commentsExtended+xml",
        )

        rels = etree.fromstring(source.read("word/_rels/document.xml.rels"))
        rel_ns = rels.tag.split("}")[0].strip("{")
        relationship = etree.SubElement(rels, f"{{{rel_ns}}}Relationship")
        relationship.set("Id", "rExistingCommentsExt")
        relationship.set(
            "Type",
            "http://schemas.microsoft.com/office/2011/relationships/commentsExtended",
        )
        relationship.set("Target", "commentsExtended.xml")

        with zipfile.ZipFile(temp, "w", zipfile.ZIP_DEFLATED) as target:
            for name in source.namelist():
                if name == "word/comments.xml":
                    data = etree.tostring(comments, xml_declaration=True, encoding="UTF-8")
                elif name == "[Content_Types].xml":
                    data = etree.tostring(content_types, xml_declaration=True, encoding="UTF-8")
                elif name == "word/_rels/document.xml.rels":
                    data = etree.tostring(rels, xml_declaration=True, encoding="UTF-8")
                else:
                    data = source.read(name)
                target.writestr(name, data)
            target.writestr(
                "word/commentsExtended.xml",
                etree.tostring(extended, xml_declaration=True, encoding="UTF-8"),
            )
    temp.replace(path)


def _set_settings_ignorable(path, value):
    temp = path.with_suffix(".settings.docx")
    with zipfile.ZipFile(path, "r") as source:
        settings = etree.fromstring(source.read("word/settings.xml"))
        settings.set(mcn("Ignorable"), value)
        with zipfile.ZipFile(temp, "w", zipfile.ZIP_DEFLATED) as target:
            for name in source.namelist():
                data = (
                    etree.tostring(settings, xml_declaration=True, encoding="UTF-8")
                    if name == "word/settings.xml"
                    else source.read(name)
                )
                target.writestr(name, data)
    temp.replace(path)


def _add_complex_field(paragraph, instruction, result):
    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(docx_qn("w:fldCharType"), "begin")
    begin_run.append(begin)
    paragraph._p.append(begin_run)

    instruction_run = OxmlElement("w:r")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.text = instruction
    instruction_run.append(instruction_text)
    paragraph._p.append(instruction_run)

    separate_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(docx_qn("w:fldCharType"), "separate")
    separate_run.append(separate)
    paragraph._p.append(separate_run)

    paragraph.add_run(result)

    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(docx_qn("w:fldCharType"), "end")
    end_run.append(end)
    paragraph._p.append(end_run)


def _make_source(path):
    doc = Document()
    p0 = doc.add_paragraph()
    prefix = p0.add_run("前缀")
    prefix.bold = True
    middle = p0.add_run("错别")
    middle.italic = True
    p0.add_run("字后缀")

    doc.add_paragraph("这是润色建议。")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run("表格里的错误。")

    p3 = doc.add_paragraph("请查看")
    _add_hyperlink(p3, "官方链接", "https://example.com")

    p4 = doc.add_paragraph()
    p4.add_run("保留旧修订")
    _wrap_last_run_in_revision(p4)
    doc.add_paragraph("重复词和重复词。")
    doc.add_comment(prefix, text="existing comment", author="Human")
    doc.save(path)
    _add_existing_comment_extension(path)


def _issues():
    return [
        {
            "fix_class": "must_fix", "location": "P0", "current": "错别字",
            "suggested": "错别词", "reason": "错别字", "category": "文字",
        },
        {
            "fix_class": "polish", "location": "P1", "current": "润色建议",
            "suggested": "表达建议", "reason": "表达可优化", "category": "润色",
        },
        {
            "fix_class": "must_fix", "location": "P2", "current": "错误",
            "suggested": "差错", "reason": "表格回写", "category": "文字",
        },
        {
            "fix_class": "must_fix", "location": "P3", "current": "官方链接",
            "suggested": "权威链接", "reason": "嵌套结构", "category": "链接",
        },
        {
            "fix_class": "must_fix", "location": "P4", "current": "旧修订",
            "suggested": "新修订", "reason": "既有修订", "category": "冲突",
        },
        {
            "fix_class": "must_fix", "location": "P5", "current": "重复词",
            "suggested": "替换词", "reason": "歧义锚点", "category": "冲突",
        },
    ]


class WordWritebackTests(unittest.TestCase):
    def test_preserves_runs_and_builds_valid_comments(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "sample.docx"
            output = root / "deliverables" / "custom-review.docx"
            _make_source(source)

            result = _run_02_writeback(
                str(source), "sample", _issues(), "Codex审校",
                out_path=str(output))

            self.assertEqual(result, str(output))
            self.assertTrue(source.exists())
            self.assertTrue(output.exists())
            Document(output)
            report = audit_docx(
                output, expected_author="Codex审校", require_revisions=True)
            self.assertEqual(report["relationships"], "ok")

            with zipfile.ZipFile(output) as package:
                names = set(package.namelist())
                self.assertIn("word/comments.xml", names)
                self.assertIn("word/commentsExtended.xml", names)
                document = etree.fromstring(package.read("word/document.xml"))
                comments = etree.fromstring(package.read("word/comments.xml"))
                comments_extended = etree.fromstring(
                    package.read("word/commentsExtended.xml"))

            refs = list(document.iter(qn("commentReference")))
            self.assertEqual(len(refs), 6)
            self.assertTrue(all(ref.getparent().tag == qn("r") for ref in refs))
            for ref in refs:
                style = ref.getparent().find(qn("rPr") + "/" + qn("rStyle"))
                self.assertIsNotNone(style)
                self.assertEqual(style.get(qn("val")), "CommentReference")
            old_ext = next(
                item for item in comments_extended.iter(f"{{{W15}}}commentEx")
                if item.get(f"{{{W15}}}paraId") == "ABCDEF01"
            )
            self.assertEqual(old_ext.get(f"{{{W15}}}done"), "1")

            codex_insertions = [
                node for node in document.iter(qn("ins"))
                if node.get(qn("author")) == "Codex审校"
            ]
            codex_deletions = [
                node for node in document.iter(qn("del"))
                if node.get(qn("author")) == "Codex审校"
            ]
            self.assertEqual(len(codex_insertions), 2)
            self.assertEqual(len(codex_deletions), 2)
            self.assertTrue(any(
                node.get(qn("author")) == "Human"
                for node in document.iter(qn("ins"))))
            self.assertTrue(any(
                node.tag == qn("hyperlink") for node in document.iter()))
            self.assertIn("官方链接", "".join(
                text.text or "" for text in document.iter(qn("t"))))

            first_para = document.find(".//" + qn("body") + "/" + qn("p"))
            accepted = "".join(
                text.text or "" for text in first_para.iter(qn("t"))
                if not any(a.tag == qn("del") for a in text.iterancestors())
            )
            self.assertEqual(accepted, "前缀错别词后缀")
            rejected_parts = []
            for node in first_para.iter():
                if node.tag == qn("delText"):
                    rejected_parts.append(node.text or "")
                elif (node.tag == qn("t")
                      and not any(a.tag == qn("ins") for a in node.iterancestors())):
                    rejected_parts.append(node.text or "")
            self.assertEqual("".join(rejected_parts), "前缀错别字后缀")
            prefix_run = next(
                run for run in first_para.iter(qn("r"))
                if "".join(text.text or "" for text in run.iter(qn("t"))) == "前缀"
            )
            self.assertIsNotNone(prefix_run.find(qn("rPr") + "/" + qn("b")))

            comment_text = "\n".join(
                "".join(text.text or "" for text in comment.iter(qn("t")))
                for comment in comments.iter(qn("comment"))
            )
            self.assertEqual(comment_text.count("自动修订已跳过"), 2)

    def test_wrong_explicit_location_never_falls_back_to_another_paragraph(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "wrong-location.docx"
            output = root / "review.docx"
            doc = Document()
            doc.add_paragraph("唯一锚点")
            doc.add_paragraph("第二段")
            doc.save(source)

            output.write_bytes(b"existing review")
            with self.assertRaisesRegex(RuntimeError, "no findings could be located"):
                _run_02_writeback(
                    str(source), "wrong-location", [{
                        "fix_class": "must_fix", "location": "P1",
                        "current": "唯一锚点", "suggested": "错误改写",
                        "reason": "位置故意写错", "category": "定位",
                    }], "Codex审校", out_path=str(output))
            self.assertEqual(output.read_bytes(), b"existing review")

    def test_tab_and_complex_field_are_comment_only_and_preserved(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "unsafe-structure.docx"
            output = root / "review.docx"
            doc = Document()
            tab_paragraph = doc.add_paragraph()
            tab_paragraph.add_run("甲")
            tab_paragraph.add_run().add_tab()
            tab_paragraph.add_run("乙")
            field_paragraph = doc.add_paragraph()
            _add_complex_field(field_paragraph, " DATE ", "旧值")
            doc.save(source)

            issues = [
                {
                    "fix_class": "must_fix", "location": "P0",
                    "current": "甲乙", "suggested": "丙丁",
                    "reason": "不得跨越制表符", "category": "结构",
                },
                {
                    "fix_class": "must_fix", "location": "P1",
                    "current": "旧值", "suggested": "新值",
                    "reason": "不得改写域结果", "category": "结构",
                },
            ]
            _run_02_writeback(
                str(source), "unsafe-structure", issues, "Codex审校",
                out_path=str(output))

            with zipfile.ZipFile(output) as package:
                document = etree.fromstring(package.read("word/document.xml"))
                comments = etree.fromstring(package.read("word/comments.xml"))
            codex_revisions = [
                node for tag in ("ins", "del") for node in document.iter(qn(tag))
                if node.get(qn("author")) == "Codex审校"
            ]
            self.assertEqual(codex_revisions, [])

            body = document.find(qn("body"))
            paragraphs = body.findall(qn("p"))
            tab_children = list(paragraphs[0])
            first_run = next(
                node for node in tab_children
                if node.tag == qn("r")
                and "".join(t.text or "" for t in node.iter(qn("t"))) == "甲"
            )
            tab_run = next(
                node for node in tab_children
                if node.tag == qn("r") and node.find(qn("tab")) is not None
            )
            last_run = next(
                node for node in tab_children
                if node.tag == qn("r")
                and "".join(t.text or "" for t in node.iter(qn("t"))) == "乙"
            )
            self.assertLess(tab_children.index(first_run), tab_children.index(tab_run))
            self.assertLess(tab_children.index(tab_run), tab_children.index(last_run))
            self.assertEqual(len(list(paragraphs[0].iter(qn("tab")))), 1)

            field_types = [
                node.get(qn("fldCharType"))
                for node in paragraphs[1].iter(qn("fldChar"))
            ]
            self.assertEqual(field_types, ["begin", "separate", "end"])
            self.assertEqual(
                "".join(node.text or "" for node in paragraphs[1].iter(qn("instrText"))),
                " DATE ",
            )
            self.assertIn("旧值", "".join(
                node.text or "" for node in paragraphs[1].iter(qn("t"))))
            self.assertNotIn("新值", "".join(
                node.text or "" for node in paragraphs[1].iter(qn("t"))))
            codex_comments = [
                node for node in comments.iter(qn("comment"))
                if node.get(qn("author")) == "Codex审校"
            ]
            self.assertEqual(len(codex_comments), 2)

    def test_short_anchor_cannot_expand_into_long_sentence(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "short-anchor.docx"
            output = root / "review.docx"
            doc = Document()
            doc.add_paragraph("错字")
            doc.save(source)
            suggestion = "这是一个被错误放大的完整句子，足以覆盖原来的两个字，而且会造成严重改写。"

            _run_02_writeback(
                str(source), "short-anchor", [{
                    "fix_class": "must_fix", "location": "P0",
                    "current": "错字", "suggested": suggestion,
                    "reason": "短锚点扩写", "category": "振幅",
                }], "Codex审校", out_path=str(output))

            with zipfile.ZipFile(output) as package:
                document = etree.fromstring(package.read("word/document.xml"))
                comments = etree.fromstring(package.read("word/comments.xml"))
            codex_revisions = [
                node for tag in ("ins", "del") for node in document.iter(qn(tag))
                if node.get(qn("author")) == "Codex审校"
            ]
            self.assertEqual(codex_revisions, [])
            self.assertEqual(Document(output).paragraphs[0].text, "错字")
            comment_text = "".join(
                node.text or "" for node in comments.iter(qn("t")))
            self.assertIn("自动修订已跳过", comment_text)
            self.assertIn(suggestion, comment_text)

    def test_settings_ignorable_tokens_are_preserved(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "settings.docx"
            output = root / "review.docx"
            doc = Document()
            doc.add_paragraph("错字")
            doc.save(source)
            _set_settings_ignorable(source, "w10 sl")

            _run_02_writeback(
                str(source), "settings", [{
                    "fix_class": "must_fix", "location": "P0",
                    "current": "错字", "suggested": "正字",
                    "reason": "字符修订", "category": "文字",
                }], "Codex审校", out_path=str(output))

            with zipfile.ZipFile(output) as package:
                settings = etree.fromstring(package.read("word/settings.xml"))
                document = etree.fromstring(package.read("word/document.xml"))
            self.assertEqual(
                set(settings.get(mcn("Ignorable"), "").split()),
                {"w10", "sl", "w14"},
            )
            self.assertTrue(any(
                node.get(qn("author")) == "Codex审校"
                for node in document.iter(qn("ins"))))

    def test_empty_issues_rejects_without_reusing_existing_output(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "empty.docx"
            output = root / "stale-review.docx"
            doc = Document()
            doc.add_paragraph("原文")
            doc.save(source)
            output.write_bytes(b"stale review")

            with self.assertRaisesRegex(ValueError, "没有可写回的 Word findings"):
                _run_02_writeback(
                    str(source), "empty", [], "Codex审校",
                    out_path=str(output))
            self.assertEqual(output.read_bytes(), b"stale review")

    def test_symlinked_output_path_cannot_overwrite_source(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            real = root / "real"
            real.mkdir()
            source = real / "source.docx"
            doc = Document()
            doc.add_paragraph("错字")
            doc.save(source)
            original = source.read_bytes()
            link = root / "linked"
            link.symlink_to(real, target_is_directory=True)

            with self.assertRaisesRegex(ValueError, "不能覆盖源 DOCX"):
                _run_02_writeback(
                    str(source), "source", [{
                        "fix_class": "must_fix", "location": "P0",
                        "current": "错字", "suggested": "正字",
                        "reason": "修订", "category": "文字",
                    }], "Codex审校", out_path=str(link / "source.docx"))
            self.assertEqual(source.read_bytes(), original)

    def test_source_change_after_snapshot_preserves_existing_output(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "changing.docx"
            output = root / "existing.docx"
            doc = Document()
            doc.add_paragraph("错字")
            doc.save(source)
            source_hash = hashlib.sha256(source.read_bytes()).hexdigest()
            output.write_bytes(b"existing review")
            original_copy = shutil.copy2

            def copy_then_change(src, dst, *args, **kwargs):
                result = original_copy(src, dst, *args, **kwargs)
                Path(src).write_bytes(Path(src).read_bytes() + b"changed")
                return result

            with patch("shutil.copy2", side_effect=copy_then_change):
                with self.assertRaisesRegex(RuntimeError, "写回期间发生变化"):
                    _run_02_writeback(
                        str(source), "changing", [{
                            "fix_class": "must_fix", "location": "P0",
                            "current": "错字", "suggested": "正字",
                            "reason": "修订", "category": "文字",
                        }], "Codex审校", out_path=str(output),
                        expected_source_sha256=source_hash)
            self.assertEqual(output.read_bytes(), b"existing review")

    def test_propagates_engine_failure(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "sample.docx"
            _make_source(source)
            existing_output = root / "failed.docx"
            existing_output.write_bytes(b"keep existing review")
            failure = SimpleNamespace(
                returncode=7, stdout="", stderr="synthetic failure")

            with patch.object(subprocess, "run", return_value=failure):
                with self.assertRaisesRegex(RuntimeError, "退出码 7"):
                    _run_02_writeback(
                        str(source), "sample", _issues()[:1], "Codex审校",
                        out_path=str(existing_output),
                    )
            self.assertEqual(existing_output.read_bytes(), b"keep existing review")

    def test_audit_parses_every_xml_part(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "sample.docx"
            broken = root / "broken.docx"
            _make_source(source)
            with zipfile.ZipFile(source, "r") as package:
                with zipfile.ZipFile(broken, "w", zipfile.ZIP_DEFLATED) as target:
                    for name in package.namelist():
                        target.writestr(
                            name,
                            b"<broken" if name == "word/styles.xml" else package.read(name),
                        )
            with self.assertRaisesRegex(DocxAuditError, "invalid XML part word/styles.xml"):
                audit_docx(broken)


if __name__ == "__main__":
    unittest.main()
