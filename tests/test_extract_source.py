import contextlib
import io
import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from src.cli import _docx_to_md, _findings_v1_to_issues, _validate_findings_source
from src.extract_source import build_review_source, extract_source, sha256_file


class ExtractSourceTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.root = Path(self.temp_dir.name)

    def tearDown(self):
        self.temp_dir.cleanup()

    def _make_docx(self) -> Path:
        path = self.root / "sample.docx"
        document = Document()
        document.add_heading("章标题", level=1)
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "表格甲"
        table.cell(0, 1).text = "表格乙"
        document.add_paragraph("正文结尾")
        document.save(path)
        return path

    def test_docx_units_and_markdown_follow_writeback_order(self):
        path = self._make_docx()

        payload = build_review_source(path)

        self.assertEqual(payload["schema"], "ai-proofread.source.v1")
        self.assertEqual(payload["source_sha256"], sha256_file(path))
        self.assertEqual(
            [(unit["location"], unit["kind"], unit["text"]) for unit in payload["units"]],
            [
                ("P0", "paragraph", "章标题"),
                ("P1", "table_cell", "表格甲"),
                ("P2", "table_cell", "表格乙"),
                ("P3", "paragraph", "正文结尾"),
            ],
        )
        self.assertEqual(payload["units"][0]["heading_level"], 1)

        md_path, md_text = _docx_to_md(path)
        self.assertEqual(md_path.read_text(encoding="utf-8"), md_text)
        self.assertEqual(md_text, "# 章标题\n表格甲\n表格乙\n正文结尾")

    def test_pdf_has_one_based_units_and_keeps_blank_pages(self):
        import fitz

        path = self.root / "sample.pdf"
        document = fitz.open()
        page = document.new_page()
        page.insert_text((72, 72), "Page one")
        document.new_page()
        document.save(path)
        document.close()

        payload = build_review_source(path)

        self.assertEqual(payload["source_type"], "pdf")
        self.assertEqual([unit["page"] for unit in payload["units"]], [1, 2])
        self.assertIn("Page one", payload["units"][0]["text"])
        self.assertTrue(payload["units"][0]["has_text"])
        self.assertFalse(payload["units"][1]["has_text"])

    def test_textless_pdf_requires_ocr(self):
        import fitz

        path = self.root / "scan.pdf"
        document = fitz.open()
        document.new_page()
        document.save(path)
        document.close()

        with self.assertRaisesRegex(RuntimeError, "OCR"):
            build_review_source(path)

    def test_cli_x_writes_source_schema(self):
        source = self._make_docx()
        output = self.root / "nested" / "review_source.json"

        completed = subprocess.run(
            [sys.executable, "-m", "src.cli", "x", str(source), "--out", str(output)],
            cwd=Path(__file__).resolve().parents[1],
            capture_output=True,
            text=True,
            check=False,
        )

        self.assertEqual(completed.returncode, 0, completed.stderr or completed.stdout)
        self.assertEqual(json.loads(output.read_text(encoding="utf-8"))["schema"],
                         "ai-proofread.source.v1")

    def test_extract_refuses_to_overwrite_source(self):
        source = self._make_docx()
        original = source.read_bytes()

        with self.assertRaisesRegex(ValueError, "不能覆盖源文件"):
            extract_source(source, source)

        self.assertEqual(source.read_bytes(), original)

    def test_extract_rejects_a_source_changed_during_snapshot(self):
        source = self._make_docx()
        with patch(
                "src.extract_source.sha256_file",
                side_effect=["1" * 64, "2" * 64]):
            with self.assertRaisesRegex(RuntimeError, "提取期间源文件发生变化"):
                build_review_source(source)

    def test_findings_v1_validates_hash_and_appends_evidence(self):
        source = self._make_docx()
        payload = {
            "schema": "ai-proofread.findings.v1",
            "source_sha256": sha256_file(source),
            "issues": [{
                "fix_class": "must_fix",
                "location": "P0",
                "current": "章标题",
                "suggested": "新标题",
                "reason": "事实修正",
                "category": "事实错误",
                "evidence": [{
                    "title": "权威页面",
                    "url": "https://example.org/fact",
                    "accessed_at": "2026-08-06",
                }],
            }],
        }

        _validate_findings_source(str(source), payload)
        issues = _findings_v1_to_issues(payload)

        self.assertEqual(issues[0]["fix_class"], "must_fix")
        self.assertIn("https://example.org/fact", issues[0]["reason"])
        with contextlib.redirect_stdout(io.StringIO()):
            with self.assertRaises(SystemExit):
                _validate_findings_source(
                    str(source), {**payload, "source_sha256": "0" * 64})

    def test_word_output_cannot_overwrite_findings(self):
        source = self._make_docx()
        findings = self.root / "findings.json"
        findings.write_text(json.dumps({
            "schema": "ai-proofread.findings.v1",
            "source_sha256": sha256_file(source),
            "issues": [{
                "fix_class": "verify",
                "location": "P0",
                "current": "章标题",
                "suggested": "",
                "reason": "Check the title.",
                "category": "title",
            }],
        }), encoding="utf-8")
        original = findings.read_bytes()
        completed = subprocess.run(
            [sys.executable, "-m", "src.cli", "w", str(source),
             "--findings", str(findings), "--out", str(findings)],
            cwd=Path(__file__).resolve().parents[1],
            capture_output=True,
            text=True,
            check=False,
        )
        self.assertNotEqual(completed.returncode, 0)
        self.assertIn("Word 输出路径不能覆盖", completed.stdout)
        self.assertEqual(findings.read_bytes(), original)

    def test_legacy_word_findings_require_source_manifest(self):
        source = self._make_docx()
        legacy = {"llm": []}
        with contextlib.redirect_stdout(io.StringIO()):
            with self.assertRaises(SystemExit):
                _validate_findings_source(str(source), legacy)

        manifest = self.root / "review_source.json"
        extract_source(source, manifest)
        _validate_findings_source(
            str(source), legacy, source_manifest=str(manifest))
        _validate_findings_source(
            str(source), [], source_manifest=str(manifest))


if __name__ == "__main__":
    unittest.main()
