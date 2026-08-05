#!/usr/bin/env python3
"""Apply proofreading findings back to the original .docx -- must_fix → track changes,
polish/verify → annotations (highlight + comment). 02-side version: pure-Chinese,
P-number located, no source-text evidence.

Modes per fix_class
-------------------
must_fix  → w:del/w:ins track changes (auto-edit) + yellow highlight + styled comment (reason/evidence chain)
polish    → yellow highlight + styled comment (suggestion only, never auto-edit)
verify    → yellow highlight + styled comment (pending verification, never auto-edit)

All comments use 01-aligned multi-run format:
  • Bold title: 【{prefix}】{category}
  • Blue (#1155CC) suggestion: ▶ 建议：{suggested}
  • Gray (#595959) evidence:   ◎ 依据：{reason}

Usage:
    python3 apply_corrections_02.py <output_dir> [--author 审阅助手] [--out <path>]

Reads:   <output_dir>/results/*.json  (issue list with fix_class)
         <output_dir> 所在目录的原始 docx
Writes:  <output_dir>/xxxx_审阅版.docx

The original docx is found by scanning the parent of output_dir for a matching
.docx file (same stem as output_dir name after "output_" prefix).
"""

import argparse
import glob
import json
import os
import posixpath
import re
import sys
import zipfile
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from lxml import etree

# ── OOXML namespaces ──────────────────────────────────────────────────────
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
MC = 'http://schemas.openxmlformats.org/markup-compatibility/2006'
CT = 'http://schemas.openxmlformats.org/package/2006/content-types'
REL  = 'http://schemas.openxmlformats.org/package/2006/relationships'
W14 = 'http://schemas.microsoft.com/office/word/2010/wordml'
W15 = 'http://schemas.microsoft.com/office/word/2012/wordml'  # WPS required for commentsExtended
NS = {'w': W, 'r': R, 'mc': MC}

# Register namespace prefixes for lxml serialization — ensures comments.xml
# uses standard w: prefix instead of auto-assigned ns0:
etree.register_namespace('w', W)
etree.register_namespace('w14', W14)
etree.register_namespace('w15', W15)
etree.register_namespace('mc', MC)
etree.register_namespace('r', R)


def qn(tag):
    return f'{{{W}}}{tag}'


def rn(tag):
    return f'{{{R}}}{tag}'


def mcn(tag):
    return f'{{{MC}}}{tag}'


# ── 振幅预算：02 无源文证据，保守设闸防 LLM 过度改写 ───────────────
MTF_CHANGED_SPAN_CAP = 30   # 原文被改最大字数
MTF_RETENTION_FLOOR = 0.35   # 字 bigram 保留下限
MTF_SHORT_RETENTION = 0.10   # 短 replacement（≤4字符）的保留下限，防过度阻挡小幅修正
MTF_SHORT_SKIP_GATE = 2      # 单字修改（changed_span ≤2）跳过振幅门禁——1字修正（如「查→茬」）bigram 保留恒为0
MTF_SHORT_ADDED_CAP = 4      # 短 replacement 最多新增字符，防短锚点被扩写成整句
MTF_ADDED_SPAN_CAP = 30      # 一次修订允许新增的最大字符数
MTF_BOUNDARY_CAP = 1         # 最大允许句界变化数
# 以上为本项目标定占位，勿跨语料照搬

# ── 锚点窗口 ──────────────────────────────────────────────────────────────
MIN_ANCHOR = 6
MAX_WINDOW = 30


def anchor_window(text, vpos, vlen):
    """Wrap seg[vpos:vpos+vlen] with ≥MIN_ANCHOR chars, unique in text."""
    s, e = vpos, vpos + vlen
    L = len(text)
    while e - s <= MAX_WINDOW:
        if e - s >= MIN_ANCHOR and text.count(text[s:e]) == 1:
            return s, e
        grew = False
        if e < L:
            e += 1
            grew = True
        if s > 0:
            s -= 1
            grew = True
        if not grew:
            break
    return None


# ── load findings ─────────────────────────────────────────────────────────
def load_findings(results_dir):
    """Load all results/*.json, keep only issues with fix_class + location + current."""
    findings = []
    for path in sorted(glob.glob(os.path.join(results_dir, '*.json'))):
        with open(path, encoding='utf-8') as f:
            data = json.load(f)
        issues = data.get('issues', []) if isinstance(data, dict) else (data if isinstance(data, list) else [])
        for issue in issues:
            fc = str(issue.get('fix_class', '')).strip()
            loc = str(issue.get('location', '')).strip()
            cur = str(issue.get('current', '')).strip()
            sug = str(issue.get('suggested', '')).strip()
            if not fc or not loc or not cur:
                continue
            pn = _extract_pn(loc)
            if pn is None:
                continue
            findings.append({
                'fix_class': fc,
                'location': loc,
                'pn': pn,
                'current': cur,
                'suggested': sug,
                'reason': str(issue.get('reason', '')).strip(),
                'description': str(issue.get('description', '')).strip(),
                'category': str(issue.get('category', '')).strip(),
            })
    return findings


def _extract_pn(loc):
    m = re.search(r'P(\d+)', loc)
    return int(m.group(1)) if m else None


# ── docx helpers (same P-numbering as extract_text.py) ────────────────────
def build_para_map(body):
    """Walk body children in order, assign P-numbers (same as extract_text.py).

    Returns {pn: w:p_element} -- only paragraphs that contain text.
    Also returns {pn: str} for text lookup.
    """
    para_map = {}
    text_map = {}
    pn = 0
    for child in body.iterchildren():
        tag = child.tag.rsplit('}', 1)[-1]
        if tag == 'p':
            t = _para_raw_text(child)
            if t.strip():
                para_map[pn] = child
                text_map[pn] = t
                pn += 1
        elif tag == 'tbl':
            for row_p in child.findall('.//' + qn('p')):
                t = _para_raw_text(row_p)
                if t.strip():
                    para_map[pn] = row_p
                    text_map[pn] = t
                    pn += 1
    return para_map, text_map


def _para_raw_text(p_elem):
    """Accept-view text of a w:p element — matches extract_text.py accepted_para_text().

    含 w:ins 内文本，排除 w:del 内文本。与 extract_text.py 的 P-number 编号保持一致，
    否则第二轮 findings 的 location:P{n} 会定位错位。
    """
    parts = []
    for t in p_elem.iter(qn('t')):
        # 排除 w:del 祖先下的被删文本
        if any(a.tag == qn('del') for a in t.iterancestors()):
            continue
        parts.append(t.text or '')
    return ''.join(parts)


def normalize_text(s):
    """标记/空白归一化：折叠连续空格，全角字母数字→半角，智能引号→直引号。

    不消除所有空格（原版消除空格后归一化位置映射回原文时偏移错误）。
    """
    # 统一非标准空格（分别替换，不拼字面序列）
    for c in (chr(0xa0), chr(0x200b), chr(0x2009), chr(0x2002), chr(0x2003)):
        s = s.replace(c, ' ')
    s = s.replace(chr(0x2013), '-').replace(chr(0x2014), '--')
    s = s.replace(chr(0x2018), "'").replace(chr(0x2019), "'")
    s = s.replace(chr(0x201c), '"').replace(chr(0x201d), '"')
    s = re.sub(r'\s+', ' ', s).strip()
    return s

def _find_para(para_map, text_map, pn, anchor):
    """Find ``anchor`` only in its declared body paragraph.

    P numbers are part of the writeback contract. Silently searching another
    paragraph can apply a valid correction to the wrong repeated sentence.
    """
    p = para_map.get(pn)
    if p is not None:
        txt = text_map.get(pn, '')
        if txt and (anchor in txt or normalize_text(anchor) in normalize_text(txt)):
            return p, txt
    return None, None


def _locate_paragraph(docx_path, p_num, body_para_map, body_text_map):
    """Locate a paragraph by P-number, checking headers/footers/textboxes if beyond body.

    Returns (lxml.Element, str) or (None, None).
    """
    # 1. Try body paragraphs first
    if p_num in body_para_map:
        return body_para_map[p_num], body_text_map.get(p_num, '')

    # 2. If beyond body, iterate headers/footers/textboxes
    body_count = len(body_para_map)
    offset = p_num - body_count
    if offset < 0:
        return None, None

    try:
        from docx import Document
        doc = Document(docx_path)
    except Exception:
        return None, None

    # Headers then footers
    for i, section in enumerate(doc.sections):
        for header in (section.header, section.first_page_header):
            if header and header.paragraphs:
                for p in header.paragraphs:
                    ht = _para_raw_text(p._element).strip()
                    if ht:
                        if offset == 0:
                            return p._element, ht
                        offset -= 1
        for footer in (section.footer, section.first_page_footer):
            if footer and footer.paragraphs:
                for p in footer.paragraphs:
                    ft = _para_raw_text(p._element).strip()
                    if ft:
                        if offset == 0:
                            return p._element, ft
                        offset -= 1

    # Textboxes
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for child in doc.element.body.iterchildren():
        for txbx in child.iter(f'{{{ns_w}}}txbxContent'):
            for p in txbx.iter(f'{{{ns_w}}}p'):
                text = _para_raw_text(p).strip()
                if text:
                    if offset == 0:
                        return p, text
                    offset -= 1

    return None, None


def bigrams(s):
    """字 bigram 计数，用于 retention 度量。"""
    s = re.sub(r'\s+', '', s or '')
    from collections import Counter
    return Counter(s[i:i + 2] for i in range(len(s) - 1)) if len(s) >= 2 else Counter()


def measure_mtf(cur, sug):
    """量化 must_fix 改动幅度：原文被改字数、保留率、句界变化。

    简化版 verify_revisions measure（不判 edit_kind，只做三指标）。
    返回值全为零 → 零动作放行。
    """
    cur = cur or ''
    sug = sug or ''
    if cur == sug:
        return {'changed_span': 0, 'added': 0, 'retention': 1.0, 'boundary_delta': 0,
                'len_cur': len(cur), 'len_sug': len(sug)}
    from difflib import SequenceMatcher
    sm = SequenceMatcher(None, cur, sug, autojunk=False)
    changed = added = 0
    retained = sum(
        i2 - i1
        for tag, i1, i2, _j1, _j2 in sm.get_opcodes()
        if tag == 'equal'
    )
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag in ('delete', 'replace'):
            changed += i2 - i1
        if tag in ('insert', 'replace'):
            added += j2 - j1
    bc = bigrams(cur)
    bs = bigrams(sug)
    inter = sum((bc & bs).values())
    retention = round(inter / sum(bc.values()), 3) if bc else 1.0
    _sent = re.compile(r'[。！？!?]+')
    bd = abs(len(_sent.findall(sug or '')) - len(_sent.findall(cur or '')))
    return {'changed_span': changed, 'added': added, 'retention': retention,
            'boundary_delta': bd, 'len_cur': len(cur), 'len_sug': len(sug)}


def _expected_revision_delta(current, suggested):
    """Return the exact deleted and inserted payload expected in OOXML."""
    from difflib import SequenceMatcher
    deleted = []
    inserted = []
    for tag, i1, i2, j1, j2 in SequenceMatcher(
            None, current, suggested, autojunk=False).get_opcodes():
        if tag in ('delete', 'replace'):
            deleted.append(current[i1:i2])
        if tag in ('insert', 'replace'):
            inserted.append(suggested[j1:j2])
    return ''.join(deleted), ''.join(inserted)


# ── Track changes helpers ─────────────────────────────────────────────────
def _next_rev_id(root):
    """Maximum existing w:revId + 1 across document.xml."""
    ids = []
    for el in root.iter(qn('ins')):
        ids.append(int(el.get(qn('id'), 0)))
    for el in root.iter(qn('del')):
        ids.append(int(el.get(qn('id'), 0)))
    return max(ids) + 1 if ids else 1


def _next_cmt_id(existing_comments):
    """Maximum comment ID + 1 from existing comments.xml."""
    ids = [0]
    if existing_comments is not None:
        for cmt in existing_comments.iter(qn('comment')):
            try:
                ids.append(int(cmt.get(qn('id'), 0)))
            except (ValueError, TypeError):
                pass
    return max(ids) + 1


def _author_iso(author):
    now = datetime.now(timezone.utc)
    return author, now.strftime('%Y-%m-%dT%H:%M:%SZ')


def _plain_run_text(run):
    """Return text for an editable plain run, otherwise ``None``.

    Tracked changes must not flatten fields, drawings, tabs, hyperlinks, or
    pre-existing revisions. Only a direct paragraph run with one text node and
    optional run properties is safe to split.
    """
    if run.tag != qn('r'):
        return None
    children = list(run)
    if any(child.tag not in (qn('rPr'), qn('t')) for child in children):
        return None
    texts = [child for child in children if child.tag == qn('t')]
    if len(texts) != 1:
        return None
    return texts[0].text or ''


def _clone_text_run(run, text, deleted=False, highlight=None):
    """Clone a plain run while replacing only its text payload."""
    clone = deepcopy(run)
    for child in list(clone):
        if child.tag != qn('rPr'):
            clone.remove(child)
    text_el = etree.SubElement(clone, qn('delText') if deleted else qn('t'))
    text_el.text = text
    text_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    rpr = clone.find(qn('rPr'))
    if highlight is not None:
        if rpr is None:
            rpr = etree.Element(qn('rPr'))
            clone.insert(0, rpr)
        old = rpr.find(qn('highlight'))
        if old is not None:
            rpr.remove(old)
        if highlight:
            etree.SubElement(rpr, qn('highlight')).set(qn('val'), highlight)
    return clone


def _safe_text_stream(para):
    """Build a text stream and direct-run intervals without crossing unsafe OOXML."""
    stream = ''
    intervals = []
    field_depth = 0
    for child in para:
        field_types = [
            node.get(qn('fldCharType'))
            for node in child.iter(qn('fldChar'))
        ]
        field_depth += sum(1 for value in field_types if value == 'begin')
        text = _plain_run_text(child)
        if text is not None and field_depth == 0:
            start = len(stream)
            stream += text
            intervals.append((start, len(stream), child, text))
        else:
            # Every non-plain child is a hard boundary, including zero-width
            # bookmarks/comment markers, field results, and runs containing
            # tabs or drawings. Never move such OOXML across a replacement.
            stream += '\ufff0'
        field_depth = max(
            0, field_depth
            - sum(1 for value in field_types if value == 'end'))
    return stream, intervals


def _locate_safe_span(para, current):
    """Locate a unique target fully contained in editable direct runs."""
    if not current:
        return None, 'empty target'
    accepted = _para_raw_text(para)
    if accepted.count(current) != 1:
        return None, 'anchor is missing or ambiguous in paragraph'
    if para.find('.//' + qn('permStart')) is not None:
        return None, 'paragraph contains a protected range'
    stream, intervals = _safe_text_stream(para)
    if stream.count(current) != 1:
        return None, 'anchor crosses a hyperlink, field, or existing revision'
    start = stream.find(current)
    end = start + len(current)
    affected = [item for item in intervals if item[0] < end and item[1] > start]
    if not affected or affected[0][0] > start or affected[-1][1] < end:
        return None, 'anchor crosses an unsafe run boundary'
    return (start, end, intervals, affected), None


def _source_run_pieces(intervals, start, end):
    pieces = []
    for run_start, run_end, run, text in intervals:
        left = max(start, run_start)
        right = min(end, run_end)
        if left < right:
            pieces.append((run, text[left - run_start:right - run_start]))
    return pieces


def _revision_element(tag, rid, author, date_iso, pieces, inserted_text=''):
    revision = etree.Element(qn(tag))
    revision.set(qn('id'), str(rid))
    revision.set(qn('author'), author)
    revision.set(qn('date'), date_iso)
    if tag == 'del':
        for run, text in pieces:
            revision.append(_clone_text_run(run, text, deleted=True, highlight=False))
    elif inserted_text:
        template = pieces[0][0]
        revision.append(_clone_text_run(
            template, inserted_text, deleted=False, highlight='yellow'))
    return revision


def apply_track_change(para, current, suggested, rid, author, date_iso):
    """Apply one tracked change while preserving all unrelated OOXML nodes."""
    located, conflict = _locate_safe_span(para, current)
    if located is None:
        print(f'  WARN: track change downgraded: {conflict}: "{current[:40]}"')
        return False
    cpos, c_end, intervals, affected = located
    first_start, _, first_run, first_text = affected[0]
    _, last_end, last_run, last_text = affected[-1]
    replacements = []

    prefix = first_text[:cpos - first_start]
    if prefix:
        replacements.append(_clone_text_run(first_run, prefix))

    from difflib import SequenceMatcher
    sm = SequenceMatcher(None, current, suggested, autojunk=False)
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        pieces = _source_run_pieces(intervals, cpos + i1, cpos + i2)
        if tag == 'equal':
            for run, text in pieces:
                replacements.append(_clone_text_run(run, text))
        elif tag == 'delete':
            replacements.append(_revision_element(
                'del', rid, author, date_iso, pieces))
        elif tag == 'insert':
            style_pieces = _source_run_pieces(
                intervals, cpos + max(0, i1 - 1), cpos + max(1, i1))
            style_pieces = style_pieces or [(first_run, '')]
            replacements.append(_revision_element(
                'ins', rid, author, date_iso, style_pieces,
                inserted_text=suggested[j1:j2]))
        elif tag == 'replace':
            replacements.append(_revision_element(
                'del', rid, author, date_iso, pieces))
            replacements.append(_revision_element(
                'ins', rid, author, date_iso, pieces or [(first_run, '')],
                inserted_text=suggested[j1:j2]))

    suffix = last_text[len(last_text) - (last_end - c_end):] if last_end > c_end else ''
    if suffix:
        replacements.append(_clone_text_run(last_run, suffix))

    insert_at = list(para).index(first_run)
    for _, _, run, _ in affected:
        para.remove(run)
    for element in replacements:
        para.insert(insert_at, element)
        insert_at += 1

    return True


def _comment_reference_run(cmt_id):
    run = etree.Element(qn('r'))
    rpr = etree.SubElement(run, qn('rPr'))
    style = etree.SubElement(rpr, qn('rStyle'))
    style.set(qn('val'), 'CommentReference')
    cref = etree.SubElement(run, qn('commentReference'))
    cref.set(qn('id'), str(cmt_id))
    return run


def _highlight_run(run):
    rpr = run.find(qn('rPr'))
    if rpr is None:
        rpr = etree.Element(qn('rPr'))
        run.insert(0, rpr)
    highlight = rpr.find(qn('highlight'))
    if highlight is None:
        highlight = etree.SubElement(rpr, qn('highlight'))
    highlight.set(qn('val'), 'yellow')


def _top_level_child(para, node):
    while node is not None and node.getparent() is not para:
        node = node.getparent()
    return node


def apply_annotation(para, current, fix_class, rid, author, date_iso, suggested='', reason='', description='', category=''):
    """Apply yellow highlight + comment to `current` in the paragraph.

    Adds highlight to runs that contain the error text, then places
    commentRangeStart before the first affected run and
    commentRangeEnd + commentReference after the last affected run.
    """
    located, _ = _locate_safe_span(para, current)
    cmt_id = rid
    if located is not None:
        cpos, c_end, intervals, affected = located
        first_start, _, first_run, first_text = affected[0]
        _, last_end, last_run, last_text = affected[-1]
        replacements = []
        prefix = first_text[:cpos - first_start]
        if prefix:
            replacements.append(_clone_text_run(first_run, prefix))
        highlighted = []
        for run, text in _source_run_pieces(intervals, cpos, c_end):
            clone = _clone_text_run(run, text, highlight='yellow')
            replacements.append(clone)
            highlighted.append(clone)
        suffix = last_text[len(last_text) - (last_end - c_end):] if last_end > c_end else ''
        if suffix:
            replacements.append(_clone_text_run(last_run, suffix))
        insert_at = list(para).index(first_run)
        for _, _, run, _ in affected:
            para.remove(run)
        for element in replacements:
            para.insert(insert_at, element)
            insert_at += 1
        first_anchor = highlighted[0]
        last_anchor = highlighted[-1]
    else:
        # Annotation-only fallback for hyperlinks, fields, protected ranges, and
        # existing revisions. The nested structure remains byte-for-byte intact.
        accepted = _para_raw_text(para)
        if accepted.count(current) != 1:
            print(f'  WARN (annotate): "{current[:40]}" missing or ambiguous')
            return False
        cpos = accepted.find(current)
        c_end = cpos + len(current)
        joined = ''
        run_info = []
        for run in para.iter(qn('r')):
            if any(a.tag == qn('del') for a in run.iterancestors()):
                continue
            text = ''.join(t.text or '' for t in run.iter(qn('t')))
            if not text:
                continue
            run_info.append((len(joined), len(joined) + len(text), run))
            joined += text
        affected_runs = [run for start, end, run in run_info
                         if start < c_end and end > cpos]
        if not affected_runs:
            return False
        for run in affected_runs:
            _highlight_run(run)
        first_anchor = _top_level_child(para, affected_runs[0])
        last_anchor = _top_level_child(para, affected_runs[-1])
        if first_anchor is None or last_anchor is None:
            return False

    crs = etree.Element(qn('commentRangeStart'))
    crs.set(qn('id'), str(cmt_id))
    para.insert(list(para).index(first_anchor), crs)
    cre = etree.Element(qn('commentRangeEnd'))
    cre.set(qn('id'), str(cmt_id))
    last_idx_in_parent = list(para).index(last_anchor)
    para.insert(last_idx_in_parent + 1, cre)
    para.insert(last_idx_in_parent + 2, _comment_reference_run(cmt_id))

    return cmt_id, category, fix_class, current, suggested, reason, description


# ── comment.xml ───────────────────────────────────────────────────────────
def build_comment_xml(comments_list):
    """Build <w:comments> from structured entries.

    Each entry: (cmt_id, author, date_iso, title, suggestion, evidence)
    Creates styled multi-run comment matching 01 pipeline format:
      - Bold title: 【prefix】category
      - Blue (#1155CC) suggestion: ▶ 建议：...
      - Gray (#595959) evidence:   ◎ 依据：...

    The root element includes mc:Ignorable="w14" and full namespace
    declarations so Word correctly renders the comments pane.
    """
    comments = etree.fromstring(f'<w:comments xmlns:w="{W}"/>')

    for entry in comments_list:
        cmt_id, author, date_iso, title, suggestion, evidence = entry
        c = etree.SubElement(comments, qn('comment'))
        c.set(qn('id'), str(cmt_id))
        c.set(qn('author'), author)
        c.set(qn('date'), date_iso)
        c.set(qn('initials'), 'AI')
        p = etree.SubElement(c, qn('p'))

        # Paragraph style required for Word comment rendering
        ppr = etree.SubElement(p, qn('pPr'))
        pstyle = etree.SubElement(ppr, qn('pStyle'))
        pstyle.set(qn('val'), 'CommentText')

        # Run 0: annotationRef (standard Word comment appearance)
        r0 = etree.SubElement(p, qn('r'))
        rpr0 = etree.SubElement(r0, qn('rPr'))
        rs0 = etree.SubElement(rpr0, qn('rStyle'))
        rs0.set(qn('val'), 'CommentReference')
        etree.SubElement(r0, qn('annotationRef'))

        # Run 1: title in bold
        r1 = etree.SubElement(p, qn('r'))
        rpr1 = etree.SubElement(r1, qn('rPr'))
        etree.SubElement(rpr1, qn('b'))
        t1 = etree.SubElement(r1, qn('t'))
        t1.text = title
        t1.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

        # Line break before suggestion/evidence section
        if suggestion or evidence:
            br1 = etree.SubElement(p, qn('r'))
            etree.SubElement(br1, qn('br'))

        if suggestion:
            r2 = etree.SubElement(p, qn('r'))
            rpr2 = etree.SubElement(r2, qn('rPr'))
            c2 = etree.SubElement(rpr2, qn('color'))
            c2.set(qn('val'), '1155CC')
            t2 = etree.SubElement(r2, qn('t'))
            t2.text = f'▶ 建议：{suggestion}'
            t2.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

        if evidence:
            if suggestion:
                br2 = etree.SubElement(p, qn('r'))
                etree.SubElement(br2, qn('br'))
            r3 = etree.SubElement(p, qn('r'))
            rpr3 = etree.SubElement(r3, qn('rPr'))
            c3 = etree.SubElement(rpr3, qn('color'))
            c3.set(qn('val'), '595959')
            t3 = etree.SubElement(r3, qn('t'))
            t3.text = f'◎ 依据：{evidence}'
            t3.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    return comments


# ── WPS 兼容：批注段落样式 + commentsExtended ────────────────────────────────

def _inject_comment_styles(styles_root):
    """注入 CommentText(段落) 和 CommentReference(字符) 样式定义。

    WPS 不似 Word 会将缺失的批注样式 fallback 到 Normal，
    缺少这两条样式会导致侧边批注完全不显示。
    """
    existing = set()
    for s in styles_root.iter(qn('style')):
        sid = s.get(qn('styleId'))
        if sid:
            existing.add(sid)

    if 'CommentText' not in existing:
        st = etree.SubElement(styles_root, qn('style'))
        st.set(qn('type'), 'paragraph')
        st.set(qn('styleId'), 'CommentText')
        n = etree.SubElement(st, qn('name'))
        n.set(qn('val'), 'Comment Text')
        b = etree.SubElement(st, qn('basedOn'))
        b.set(qn('val'), 'Normal')
        ppr = etree.SubElement(st, qn('pPr'))
        sp = etree.SubElement(ppr, qn('spacing'))
        sp.set(qn('after'), '200')

    if 'CommentReference' not in existing:
        st = etree.SubElement(styles_root, qn('style'))
        st.set(qn('type'), 'character')
        st.set(qn('styleId'), 'CommentReference')
        n = etree.SubElement(st, qn('name'))
        n.set(qn('val'), 'Comment Reference')

    return styles_root


def _ensure_comment_para_ids(comments_xml):
    """Ensure every comment has a paragraph id used by commentsExtended.xml."""
    used = {
        value.upper()
        for p in comments_xml.iter(qn('p'))
        for value in [p.get(f'{{{W14}}}paraId')]
        if value
    }
    mapping = []
    next_id = 1
    for comment in comments_xml.iter(qn('comment')):
        para = comment.find(qn('p'))
        if para is None:
            para = etree.SubElement(comment, qn('p'))
        para_id = para.get(f'{{{W14}}}paraId')
        if not para_id:
            while f'{next_id:08X}' in used:
                next_id += 1
            para_id = f'{next_id:08X}'
            next_id += 1
            para.set(f'{{{W14}}}paraId', para_id)
            para.set(f'{{{W14}}}textId', '77777777')
            used.add(para_id)
        mapping.append(para_id)
    ignorable = comments_xml.get(mcn('Ignorable'), '').split()
    if 'w14' not in ignorable:
        comments_xml.set(mcn('Ignorable'), ' '.join(ignorable + ['w14']).strip())
    return mapping


def _build_comments_extended(comment_para_ids):
    """构建 <w15:commentsEx> — WPS 通过此文件将批注锚定到段落。

    标准批注文件使用 w15: namespace (Word 2012/2013+) 而非 w14:。
    每个 <w15:commentEx> 关联批注所在段落的 w14:paraId。
    不含 paraIdCommentId（标准版文件无此属性）。
    """
    root = etree.Element(f'{{{W15}}}commentsEx')
    for para_id in comment_para_ids:
        ex = etree.SubElement(root, f'{{{W15}}}commentEx')
        ex.set(f'{{{W15}}}paraId', para_id)
        ex.set(f'{{{W15}}}done', '0')
    return root


# ── altChunk 转换 ────────────────────────────────────────────────────────

class DocxAuditError(ValueError):
    """Raised when a generated DOCX package is not internally consistent."""


def _relationship_part(rel_path):
    if rel_path == '_rels/.rels':
        return ''
    marker = '/_rels/'
    if marker not in rel_path or not rel_path.endswith('.rels'):
        return ''
    prefix, name = rel_path.split(marker, 1)
    return posixpath.join(prefix, name[:-5])


def audit_docx(path, expected_comment_ids=None, expected_author=None,
               require_revisions=False, expected_revisions=None):
    """Audit the package graph and OOXML emitted by this engine."""
    expected_comment_ids = {str(i) for i in (expected_comment_ids or [])}
    expected_revisions = expected_revisions or []
    errors = []
    with zipfile.ZipFile(path, 'r') as package:
        bad = package.testzip()
        if bad:
            errors.append(f'corrupt ZIP member: {bad}')
        names = set(package.namelist())
        if 'word/document.xml' not in names:
            raise DocxAuditError('missing word/document.xml')

        for xml_path in (name for name in names
                         if name.endswith('.xml') or name.endswith('.rels')):
            try:
                etree.fromstring(package.read(xml_path))
            except etree.XMLSyntaxError as exc:
                errors.append(f'invalid XML part {xml_path}: {exc}')

        for rel_path in (name for name in names if name.endswith('.rels')):
            try:
                rels = etree.fromstring(package.read(rel_path))
            except etree.XMLSyntaxError as exc:
                errors.append(f'invalid relationships XML {rel_path}: {exc}')
                continue
            owner = _relationship_part(rel_path)
            base = posixpath.dirname(owner)
            for rel in rels:
                target = rel.get('Target', '')
                if (rel.get('TargetMode') == 'External' or not target
                        or '://' in target):
                    continue
                resolved = (target.lstrip('/') if target.startswith('/') else
                            posixpath.normpath(posixpath.join(base, target)))
                if resolved not in names:
                    errors.append(f'dangling relationship {rel_path} -> {target}')

        content_types = etree.fromstring(package.read('[Content_Types].xml'))
        overrides = {child.get('PartName') for child in content_types
                     if child.tag.endswith('Override')}
        document_targets = set()
        if 'word/_rels/document.xml.rels' in names:
            document_rels = etree.fromstring(
                package.read('word/_rels/document.xml.rels'))
            for rel in document_rels:
                if rel.get('TargetMode') == 'External':
                    continue
                target = rel.get('Target', '')
                document_targets.add(
                    target.lstrip('/') if target.startswith('/') else
                    posixpath.normpath(posixpath.join('word', target)))
        for part in ('word/comments.xml', 'word/commentsExtended.xml'):
            if part not in names:
                continue
            if f'/{part}' not in overrides:
                errors.append(f'missing Content Type override for {part}')
            if part not in document_targets:
                errors.append(f'missing document relationship for {part}')

        doc = etree.fromstring(package.read('word/document.xml'))
        comments = None
        comment_ids = set()
        if 'word/comments.xml' in names:
            comments = etree.fromstring(package.read('word/comments.xml'))
            comment_ids = {
                str(comment.get(qn('id')))
                for comment in comments.iter(qn('comment'))
            }

        starts = [str(node.get(qn('id'))) for node in doc.iter(qn('commentRangeStart'))]
        ends = [str(node.get(qn('id'))) for node in doc.iter(qn('commentRangeEnd'))]
        refs = list(doc.iter(qn('commentReference')))
        ref_ids = [str(node.get(qn('id'))) for node in refs]
        for cmt_id in expected_comment_ids:
            if starts.count(cmt_id) != 1 or ends.count(cmt_id) != 1 or ref_ids.count(cmt_id) != 1:
                errors.append(f'unpaired comment anchor id={cmt_id}')
            if cmt_id not in comment_ids:
                errors.append(f'missing comment body id={cmt_id}')
        for ref in refs:
            if str(ref.get(qn('id'))) not in expected_comment_ids:
                continue
            run = ref.getparent()
            style = run.find(qn('rPr') + '/' + qn('rStyle')) if run is not None else None
            if (run is None or run.tag != qn('r') or style is None
                    or style.get(qn('val')) != 'CommentReference'):
                errors.append(f'commentReference id={ref.get(qn("id"))} is not in a styled run')

        if 'word/commentsExtended.xml' in names:
            extended = etree.fromstring(package.read('word/commentsExtended.xml'))
            ext_ids = {
                item.get(f'{{{W15}}}paraId')
                for item in extended.iter(f'{{{W15}}}commentEx')
            }
            para_ids = {
                p.get(f'{{{W14}}}paraId')
                for p in (comments.iter(qn('p')) if comments is not None else [])
            }
            if any(value and value not in para_ids for value in ext_ids):
                errors.append('commentsExtended references unknown comment paragraphs')
            expected_para_ids = {
                comment.find(qn('p')).get(f'{{{W14}}}paraId')
                for comment in comments.iter(qn('comment'))
                if (str(comment.get(qn('id'))) in expected_comment_ids
                    and comment.find(qn('p')) is not None)
            } if comments is not None else set()
            if any(value not in ext_ids for value in expected_para_ids if value):
                errors.append('commentsExtended is missing a generated comment paragraph')

        revisions = list(doc.iter(qn('ins'))) + list(doc.iter(qn('del')))
        authored = [node for node in revisions
                    if not expected_author or node.get(qn('author')) == expected_author]
        if require_revisions and not authored:
            errors.append('expected tracked revisions were not written')
        for node in authored:
            if not node.get(qn('id')) or not node.get(qn('author')) or not node.get(qn('date')):
                errors.append('tracked revision is missing id/author/date metadata')
        for rev_id, expected_deleted, expected_inserted in expected_revisions:
            deleted = ''.join(
                text.text or ''
                for node in doc.iter(qn('del'))
                if str(node.get(qn('id'))) == str(rev_id)
                for text in node.iter(qn('delText'))
            )
            inserted = ''.join(
                text.text or ''
                for node in doc.iter(qn('ins'))
                if str(node.get(qn('id'))) == str(rev_id)
                for text in node.iter(qn('t'))
            )
            if deleted != expected_deleted or inserted != expected_inserted:
                errors.append(
                    f'revision text mismatch id={rev_id}: '
                    f'del={deleted!r}/{expected_deleted!r}, '
                    f'ins={inserted!r}/{expected_inserted!r}')

    if errors:
        raise DocxAuditError('; '.join(errors))
    return {
        'comments': len(expected_comment_ids),
        'revisions': len(authored),
        'relationships': 'ok',
    }


def _next_relationship_id(rels_xml, prefix):
    used = {rel.get('Id') for rel in rels_xml}
    if prefix not in used:
        return prefix
    number = 2
    while f'{prefix}{number}' in used:
        number += 1
    return f'{prefix}{number}'


def _convert_altchunk_to_paras(z, docx_path, doc_xml, body):
    """Convert altChunk (MHT/HTML embedded content) to standard w:p paragraphs.

    某些 PDF 转换工具/协作编辑器会产出 altChunk 格式的 docx，其正文不是
    标准的 w:p 而是嵌入的 MHT 文件。此函数读取 MHT，解析为纯文本段落，
    创建对应的 w:p/w:r/w:t XML 元素替换 body 中的 altChunk。
    """
    # Get relationships to map rId → target path
    rels_xml = z.read('word/_rels/document.xml.rels').decode('utf-8')
    rmap = {}
    for m in re.finditer(r'<Relationship[^>]*>', rels_xml):
        rid_m = re.search(r'Id="([^"]+)"', m.group())
        tgt_m = re.search(r'Target="([^"]+)"', m.group())
        if rid_m and tgt_m:
            rmap[rid_m.group(1)] = tgt_m.group(1)

    # Collect altChunk ids and their references
    ac_refs = []
    for ac in body.findall(qn('altChunk')):
        rid = ac.get(rn('id')) or ac.get(f'{{{R}}}id')
        if rid:
            target = rmap.get(rid, '')
            if target.startswith('/'):
                target = target.lstrip('/')
            elif not target.startswith('word/'):
                target = os.path.join('word', target)
            ac_refs.append((ac, rid, target))

    if not ac_refs:
        return

    # Collect all paragraph content from all altChunks
    all_lines = []
    for _, _, target in ac_refs:
        try:
            raw = z.read(target).decode('utf-8', errors='replace')
        except KeyError:
            continue

        # Strip MIME headers before first <html
        idx = raw.find('<html')
        if idx >= 0:
            raw = raw[idx:]

        # Strip HTML tags
        html = re.sub(r'<script[^>]*>.*?</script>', '', raw,
                       flags=re.DOTALL | re.IGNORECASE)
        html = re.sub(r'<style[^>]*>.*?</style>', '', html,
                       flags=re.DOTALL | re.IGNORECASE)
        html = re.sub(r'<br\s*/?>', '\n', html, flags=re.IGNORECASE)
        html = re.sub(r'</p>', '\n', html, flags=re.IGNORECASE)
        html = re.sub(r'</div>', '\n', html, flags=re.IGNORECASE)
        html = re.sub(r'<[^>]+>', '', html)
        html = re.sub(r'&nbsp;', ' ', html)
        html = re.sub(r'&amp;', '&', html)
        html = re.sub(r'&lt;', '<', html)
        html = re.sub(r'&gt;', '>', html)
        html = re.sub(r'&quot;', '"', html)

        for line in html.split('\n'):
            line = line.strip().replace('\r', '')
            if line:
                all_lines.append(line)

    if not all_lines:
        print('WARN: altChunk 提取为空', file=sys.stderr)
        return

    # Remove altChunk elements from body, keeping only sectPr
    for ac, _, _ in ac_refs:
        body.remove(ac)

    # Insert w:p elements for each text line
    for line in all_lines:
        p = etree.SubElement(body, qn('p'))  # inserted at end, before sectPr
        r = etree.SubElement(p, qn('r'))
        t = etree.SubElement(r, qn('t'))
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = line

    # Move sectPr to end of body (it was likely after the altChunks)
    sectPrs = body.findall(qn('sectPr'))
    for sp in sectPrs:
        body.remove(sp)
    if sectPrs:
        body.append(sectPrs[-1])

    print(f'  altChunk → {len(all_lines)} w:p paragraphs')


# ── main ──────────────────────────────────────────────────────────────────
def main():
    ap = argparse.ArgumentParser(description='02: 审校 findings → docx 回写')
    ap.add_argument('output_dir', help='output_{docname} 目录')
    ap.add_argument('--author', default='审阅助手')
    ap.add_argument('--out', help='输出路径（默认 output_dir/xxx_审阅版.docx）')
    args = ap.parse_args()

    d = args.output_dir
    if not os.path.isdir(d):
        print(f'ERROR: output_dir not found: {d}')
        sys.exit(1)

    # Find original docx
    parent = os.path.dirname(d) or '.'
    stem = os.path.basename(d).removeprefix('output_')
    candidates = [os.path.join(parent, f) for f in os.listdir(parent)
                  if f.startswith(stem) and f.lower().endswith('.docx')
                  and '审阅版' not in f and '修订' not in f
                  and '批注' not in f and '注释' not in f
                  and '参考' not in f and '原始备份' not in f and 'flatten' not in f
                  and not any(t in f for t in ('-批注', '-参考', '标准版', '修改版'))]
    # 选长度最短且排序最前的原始文件（排除附加词汇后的变体）
    candidates.sort(key=lambda x: (len(os.path.basename(x)), x))
    if not candidates:
        print(f'ERROR: no source DOCX matching {stem!r} in {parent}', file=sys.stderr)
        sys.exit(1)
    docx_path = candidates[0]
    # 记录选中文件用于诊断
    print(f'源文档: {os.path.basename(docx_path)}')

    # Load findings
    results_dir = os.path.join(d, 'results')
    if not os.path.isdir(results_dir):
        print(f'ERROR: results/ not found: {results_dir}')
        sys.exit(1)
    findings = load_findings(results_dir)
    must_fix = [f for f in findings if f['fix_class'] == 'must_fix']
    annotations = [f for f in findings if f['fix_class'] in ('polish', 'verify')]
    print(f'Load: {len(findings)} findings: {len(must_fix)} must_fix + {len(annotations)} polish/verify')

    if not must_fix and not annotations:
        print('No actionable findings. Skip.')
        return

    # Open docx
    try:
        z = zipfile.ZipFile(docx_path, 'r')
    except Exception as e:
        print(f'ERROR: cannot open {docx_path}: {e}')
        sys.exit(1)

    # Read document.xml
    doc_xml = etree.fromstring(z.read('word/document.xml'))
    body = doc_xml.find(qn('body'))
    para_map, text_map = build_para_map(body)
    print(f'Doc: {docx_path} → {len(para_map)} paragraphs')

    # ── Handle altChunk (embedded MHT/HTML) ──
    if not para_map:
        ac = body.findall(qn('altChunk'))
        if ac:
            print('WARN: 文档为 altChunk 格式，正在转换为标准 w:p 段落...')
            _convert_altchunk_to_paras(z, docx_path, doc_xml, body)
            para_map, text_map = build_para_map(body)
            print(f'  转换后: {len(para_map)} paragraphs')

    # Read settings.xml for trackChanges injection
    settings_xml = None
    if must_fix:
        try:
            settings_xml = etree.fromstring(z.read('word/settings.xml'))
        except KeyError:
            print('WARN: no word/settings.xml, creating skeleton')
            settings_xml = etree.Element(qn('settings'))

    # Read existing comments.xml for comment ID counter
    existing_comments_xml = None
    try:
        existing_comments_xml = etree.fromstring(z.read('word/comments.xml'))
    except (KeyError, etree.XMLSyntaxError):
        pass

    # Start separate ID counters for revisions and comments (avoid collision)
    next_rid = _next_rev_id(doc_xml)
    next_cmt_id = _next_cmt_id(existing_comments_xml)
    author, date_iso = _author_iso(args.author)

    # ── shared comment entries pool (both must_fix and annotations append here) ──
    COMMENT_ENTRIES = []  # (cmt_id, author, date_iso, title, suggestion, evidence)
    EXPECTED_REVISIONS = []  # (revision id, deleted payload, inserted payload)

    # ── apply must_fix: annotation first, then track change (with amplitude gate) ──
    # 批次处理：先定位所有 findings 的段落和位置，按段落后→前排序分批
    # 避免前一批的 run 删除破坏后一批的定位
    track_done = 0
    track_fail = 0
    downgraded = 0       # 超预算降为批注的 must_fix
    track_skip = 0       # budget check 跳过的
    must_ann_done = 0
    must_ann_fail = 0

    # Locate against the unmodified accepted view, then process each paragraph
    # from back to front so a later edit cannot shift an earlier anchor.
    pre_located = []  # [(p_el, cpos, current, suggested, fix_class, ...)]
    for fd in must_fix:
        p, txt = _find_para(para_map, text_map, fd['pn'], fd['current'])
        if p is None:
            p, txt = _locate_paragraph(docx_path, fd['pn'], para_map, text_map)
        if p is None:
            track_fail += 1
            print(f'  WARN: anchor "{fd["current"][:30]}" not found (even after full search)')
            continue

        accepted = _para_raw_text(p)
        if accepted.count(fd['current']) != 1:
            track_fail += 1
            print(f'  WARN: anchor "{fd["current"][:30]}" missing or ambiguous')
            continue
        cpos = accepted.find(fd['current'])

        pre_located.append((p, cpos, fd['current'], fd['suggested'],
                           fd.get('fix_class', 'must_fix'),
                           fd.get('reason', ''), fd.get('category', ''),
                           fd.get('description', ''), fd.get('pn', '')))

    # 按段落分组，段内从后往前排
    para_groups = {}
    for pl in pre_located:
        pid = id(pl[0])  # 用段落元素 id 分组
        para_groups.setdefault(pid, {'para': pl[0], 'findings': []})
        para_groups[pid]['findings'].append(pl)

    # Apply each finding locally; never rebuild the paragraph's complete run
    # list. Unsafe OOXML targets remain untouched and receive a comment only.
    for g in para_groups.values():
        g['findings'].sort(key=lambda x: x[1], reverse=True)
        for pl in g['findings']:
            p_el, cpos, cur, sug, fc, reason, cat, desc, pn = pl
            met = measure_mtf(cur, sug)
            min_ret = MTF_SHORT_RETENTION if met['changed_span'] <= 4 else MTF_RETENTION_FLOOR
            budget_ok = (
                (
                    met['changed_span'] <= MTF_SHORT_SKIP_GATE
                    and met['added'] <= MTF_SHORT_ADDED_CAP
                    and abs(met['len_sug'] - met['len_cur']) <= MTF_SHORT_ADDED_CAP
                )
                or (
                    met['changed_span'] <= MTF_CHANGED_SPAN_CAP
                    and met['added'] <= MTF_ADDED_SPAN_CAP
                    and met['retention'] >= min_ret
                    and met['boundary_delta'] <= MTF_BOUNDARY_CAP
                )
            )
            _, conflict = _locate_safe_span(p_el, cur)
            will_track = budget_ok and conflict is None and cur != sug
            downgrade_reason = ''
            if not budget_ok:
                downgrade_reason = (
                    f'改动超幅（改{met["changed_span"]}字/保留{met["retention"]}'
                    f'/句界{met["boundary_delta"]}）')
            elif conflict:
                downgrade_reason = conflict
            elif cur == sug:
                downgrade_reason = '建议文本与原文相同'

            ann_result = apply_annotation(
                p_el, cur, fc,
                next_cmt_id, author, date_iso,
                suggested=sug, reason=reason,
                description=desc, category=cat)
            if ann_result is False:
                must_ann_fail += 1
                continue
            cmt_id, ann_cat, ann_fc, ann_cur, ann_sug, ann_reason, ann_desc = ann_result
            must_ann_done += 1
            tracked = False
            if will_track:
                tracked = apply_track_change(
                    p_el, cur, sug, next_rid, author, date_iso)
            if tracked:
                track_done += 1
                deleted, inserted = _expected_revision_delta(cur, sug)
                EXPECTED_REVISIONS.append(
                    (str(next_rid), deleted, inserted))
                next_rid += 1
                title = f'【必改】{ann_cat}'
            else:
                downgraded += 1
                track_skip += 1
                if will_track:
                    track_fail += 1
                    downgrade_reason = '局部 OOXML 写回冲突'
                print(f'  ⛔降级: P{pn} {downgrade_reason} → 仅批注不自动改')
                title = f'待核【{ann_cat}】'
                ann_reason = '\n'.join(filter(None, [
                    ann_reason,
                    f'自动修订已跳过：{downgrade_reason}',
                ]))
            COMMENT_ENTRIES.append(
                (cmt_id, author, date_iso, title, ann_sug, ann_reason))
            next_cmt_id += 1

    # ── inject trackRevisions in settings.xml ──
    if track_done and settings_xml is not None:
        tr = settings_xml.find(qn('trackRevisions'))
        if tr is None:
            settings_xml.append(etree.Element(qn('trackRevisions')))
        ignorable = settings_xml.get(mcn('Ignorable'), '').split()
        if 'w14' not in ignorable:
            settings_xml.set(
                mcn('Ignorable'), ' '.join([*ignorable, 'w14']).strip())

    # ── apply annotations (polish/verify) ──
    annotate_done = 0
    annotate_fail = 0
    # COMMENT_ENTRIES already initialized in must_fix section
    fc_prefix = {'polish': '💬润色', 'verify': '💬待核'}
    for fd in annotations:
        p, txt = _find_para(para_map, text_map, fd['pn'], fd['current'])
        if p is None:
            p, txt = _locate_paragraph(docx_path, fd['pn'], para_map, text_map)
        if p is None:
            annotate_fail += 1
            print(f'  WARN (annotate): anchor "{fd["current"][:30]}" not found')
            continue
        result = apply_annotation(p, fd['current'], fd['fix_class'],
                                  next_cmt_id, author, date_iso,
                                  suggested=fd.get('suggested', ''),
                                  reason=fd.get('reason', ''),
                                  description=fd.get('description', ''),
                                  category=fd.get('category', ''))
        if result is False:
            annotate_fail += 1
            continue
        cmt_id, cat, fc, cur, sug, reason, desc = result
        prefix = fc_prefix.get(fc, '待核')
        title = f'{prefix}【{cat}】'
        COMMENT_ENTRIES.append((cmt_id, author, date_iso, title, sug, reason))
        annotate_done += 1
        next_cmt_id += 1

    if not COMMENT_ENTRIES and track_done == 0:
        z.close()
        print(
            'ERROR: no findings could be located or annotated; output not written',
            file=sys.stderr)
        sys.exit(1)

    # ── update comments.xml ──
    comments_ext_xml = None
    if COMMENT_ENTRIES:
        try:
            old_root = etree.fromstring(z.read('word/comments.xml'))
        except (KeyError, etree.XMLSyntaxError):
            old_root = etree.fromstring(f'<w:comments xmlns:w="{W}"/>')
        new_comments = build_comment_xml(COMMENT_ENTRIES)
        for child in new_comments:
            old_root.append(child)
        comment_para_ids = _ensure_comment_para_ids(old_root)
        try:
            comments_ext_xml = etree.fromstring(
                z.read('word/commentsExtended.xml'))
        except (KeyError, etree.XMLSyntaxError):
            comments_ext_xml = _build_comments_extended([])
        existing_ext_ids = {
            item.get(f'{{{W15}}}paraId')
            for item in comments_ext_xml.iter(f'{{{W15}}}commentEx')
        }
        for para_id in comment_para_ids:
            if para_id in existing_ext_ids:
                continue
            item = etree.SubElement(comments_ext_xml, f'{{{W15}}}commentEx')
            item.set(f'{{{W15}}}paraId', para_id)
            item.set(f'{{{W15}}}done', '0')
            existing_ext_ids.add(para_id)

    # ── write back ──
    out_path = args.out or os.path.join(d, f'{stem}_审阅版.docx')
    out_path = os.path.abspath(out_path)
    if Path(docx_path).resolve() == Path(out_path).resolve():
        print('ERROR: output path must not overwrite the source document', file=sys.stderr)
        sys.exit(1)
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    tmp_path = out_path + '.tmp'
    with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as out:
        for name in z.namelist():
            if name == 'word/document.xml':
                out.writestr(name, etree.tostring(doc_xml, xml_declaration=True, encoding='UTF-8', standalone=True))
            elif name == 'word/comments.xml' and COMMENT_ENTRIES:
                out.writestr(name, etree.tostring(old_root, xml_declaration=True, encoding='UTF-8', standalone=True))
            elif name == 'word/settings.xml' and track_done and settings_xml is not None:
                out.writestr(name, etree.tostring(settings_xml, xml_declaration=True, encoding='UTF-8', standalone=True))
            elif name == 'word/commentsExtended.xml' and comments_ext_xml is not None:
                out.writestr(name, etree.tostring(comments_ext_xml, xml_declaration=True, encoding='UTF-8', standalone=True))
            elif name == 'word/styles.xml' and COMMENT_ENTRIES:
                styles_el = etree.fromstring(z.read(name))
                _inject_comment_styles(styles_el)
                out.writestr(name, etree.tostring(styles_el, xml_declaration=True, encoding='UTF-8', standalone=True))
            else:
                out.writestr(name, z.read(name))

        # Create comments.xml if didn't exist
        if COMMENT_ENTRIES and 'word/comments.xml' not in z.namelist():
            out.writestr('word/comments.xml',
                         etree.tostring(old_root, xml_declaration=True, encoding='UTF-8', standalone=True))

        if comments_ext_xml is not None and 'word/commentsExtended.xml' not in z.namelist():
            out.writestr('word/commentsExtended.xml', etree.tostring(
                comments_ext_xml, xml_declaration=True, encoding='UTF-8', standalone=True))
        if (track_done and settings_xml is not None
                and 'word/settings.xml' not in z.namelist()):
            out.writestr('word/settings.xml', etree.tostring(
                settings_xml, xml_declaration=True, encoding='UTF-8', standalone=True))

    z.close()

    # Register only parts that were actually written. This prevents the
    # commentsExtended dangling relationship that made python-docx reject the
    # previous output.
    if COMMENT_ENTRIES:
        comments_ct = 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'
        with zipfile.ZipFile(tmp_path, 'r') as ztmp:
            package_names = set(ztmp.namelist())
            ct_xml = etree.fromstring(ztmp.read('[Content_Types].xml'))
            rels_xml = etree.fromstring(ztmp.read('word/_rels/document.xml.rels'))
        ct_ns = ct_xml.tag.split('}')[0].strip('{') if '}' in ct_xml.tag else ''
        part_specs = [
            ('word/comments.xml', comments_ct,
             'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments',
             'comments.xml', 'rComments'),
            ('word/commentsExtended.xml',
             'application/vnd.openxmlformats-officedocument.wordprocessingml.commentsExtended+xml',
             'http://schemas.microsoft.com/office/2011/relationships/commentsExtended',
             'commentsExtended.xml', 'rCommentsExt'),
        ]
        for part, content_type, rel_type, target, rel_prefix in part_specs:
            if part not in package_names:
                continue
            if not any(child.tag.endswith('Override')
                       and child.get('PartName') == f'/{part}' for child in ct_xml):
                override = etree.SubElement(ct_xml, f'{{{ct_ns}}}Override')
                override.set('PartName', f'/{part}')
                override.set('ContentType', content_type)
            if not any(rel.get('Type') == rel_type for rel in rels_xml):
                rel = etree.SubElement(rels_xml, f'{{{REL}}}Relationship')
                rel.set('Id', _next_relationship_id(rels_xml, rel_prefix))
                rel.set('Type', rel_type)
                rel.set('Target', target)

        tmp2 = tmp_path + '.rels'
        with zipfile.ZipFile(tmp_path, 'r') as zin:
            with zipfile.ZipFile(tmp2, 'w', zipfile.ZIP_DEFLATED) as zout:
                for name in zin.namelist():
                    if name == '[Content_Types].xml':
                        zout.writestr(name, etree.tostring(
                            ct_xml, xml_declaration=True, encoding='UTF-8', standalone=True))
                    elif name == 'word/_rels/document.xml.rels':
                        zout.writestr(name, etree.tostring(
                            rels_xml, xml_declaration=True, encoding='UTF-8', standalone=True))
                    else:
                        zout.writestr(name, zin.read(name))
        os.replace(tmp2, tmp_path)

    expected_comment_ids = [entry[0] for entry in COMMENT_ENTRIES]
    try:
        audit = audit_docx(
            tmp_path,
            expected_comment_ids=expected_comment_ids,
            expected_author=author,
            require_revisions=track_done > 0,
            expected_revisions=EXPECTED_REVISIONS,
        )
    except (OSError, zipfile.BadZipFile, etree.XMLSyntaxError, DocxAuditError) as exc:
        try:
            os.remove(tmp_path)
        except OSError:
            pass
        print(f'ERROR: generated DOCX failed OOXML audit: {exc}', file=sys.stderr)
        sys.exit(1)
    os.replace(tmp_path, out_path)

    print(f'\nDone: {track_done} track, {track_fail} fail, {must_ann_done} must-ann, {must_ann_fail} must-ann-fail, {annotate_done} polish/verify-ann, {annotate_fail} polish/verify-ann-fail'
          + (f' | ⛔超幅降级批注 {downgraded}' if downgraded else ''))
    print(f'Audit: {audit["comments"]} comments, {audit["revisions"]} revision nodes, relationships {audit["relationships"]}')
    print(f'→ {out_path}')


if __name__ == '__main__':
    main()
